'''
Descripttion: 
Author: wwt
Date: 2021-07-15 09:18:43
LastEditors: wwt
LastEditTime: 2021-07-22 13:24:20
'''
import os
import re
import json
import docx
import glob
import time  
import datetime
import comtypes.client
from bs4 import BeautifulSoup
from PyPDF2 import PdfFileReader
from shutil import copyfile, copytree,rmtree

SUCCESS_FLAG = "OK"
FAILED_FLAG = "NG"

# =========== 16S 分析 基类================
# 
#  BaseHTML 定义共通化方法
# ==============================================
class BaseHTML():

    def __init__(self,jsonPath=None):
        # 初始化信息
        if jsonPath:
            self.jsonData = self.readJsonFile(jsonPath)

        # 报告日期
        self.yesterda = self.getYesterday()

    '''基类 定义共通化方法'''

    # 读取json文件
    def readJsonFile(self,path):
        ''' 读取json文件 '''
        if not isinstance(path,str):
            raise TypeError("path  Parameter type error , must be str")
        
        with open(path,"r",encoding="utf-8")as f:
            data = json.load(f)
        return data

    # 删除元素
    def delTag(self,tagName,**attrs):
        ''' 删除元素 '''
        id_ = attrs.get("id",None)

        if id_:
            tdTag = self.soup.find(tagName,id=id_,**attrs)
        else:
            tdTag = self.soup.find(tagName,attrs=attrs)
        if tdTag:
            tdTag.decompose()

    # 替换文本
    def replaceText(self,name,replaceName):
        '''替换文本'''
        texts = re.sub(name,replaceName,str(self.soup))
        self.soup = texts

    # 获取前一天日期
    def getYesterday(self):
        '''获取前一天日期 ''' 
        today=datetime.date.today() 
        oneday=datetime.timedelta(days=1) 
        yesterday=today-oneday  
        return str(yesterday)

    # 写入新的文件
    def writeFile(self):
        '''写入新的文件'''
        fileName = re.search("\w+\.html",self.htmlPath)
        fileName = fileName.group()
        os.remove(self.htmlPath) 
        
        with open(self.htmlPath,"w+",encoding="utf-8") as f:
            f.write(str(self.soup))

    # WORD TO PDF
    def wordToPdf(self,sourceFile ,outFile):
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = 0
        newpdf = word.Documents.Open(sourceFile)

        try:
            if os.path.exists(outFile):
                os.remove(outFile)
            with open(outFile, 'w')as f: 
                f.close()
            newpdf.SaveAs(outFile, FileFormat=17)

            newpdf.Close()
            return SUCCESS_FLAG

        except Exception as e:
            newpdf.Close()
            return FAILED_FLAG

    # 修改 基础信息
    def editProinfo(self):

        ProjectName = self.jsonData.get("ProjectName","null")      # 项目名称
        ItemNo = self.jsonData.get("ItemNo","null")                # 子项目编号
        CustomerName = self.jsonData.get("CustomerName","null")    # 客户姓名 
        UnitAddress =  self.jsonData.get("UnitAddress","null")     # 单位 
        Received =  self.jsonData.get("Received","null")           # 收样人
        ReceivedDate =  self.jsonData.get("ReceivedDate","null")   # 收样日期
        AuditDate =  self.yesterda                                 # 审核日期 == 报告日期

        # 1.报告时间修改。
        soup = re.sub('<p align="right"><b>报告时间：.*</b></p>','<p align="right"><b>报告时间：%s</b></p>'%self.yesterda,str(self.soup))

        # 2. 删除订单编号
        tem ='''<tr>\n*\s*<td><p align="center">订单编号 </p></td>\n*\s*<td><p align="center">.*</p></td>\n*\s*<td><p align="center">销售代表 </p></td>\n*\s*<td><p align="center">.*</p></td>\n*\s*</tr>'''
        soup = re.sub(tem,"",soup)

        # 3 设置 项目名称 和 子项目编号
        projectTem = '''<td><p align="center">项目名称 </p></td>\n*\s*<td><p align="center">.*</p></td>\n*\s*<td><p align="center">子项目编号 </p></td>\n*\s*<td><p align="center">.*</p></td>'''
        projectTag = re.search(projectTem,soup)
        if  projectTag:
            replaceProjectTem = '''<td><p align="center">项目名称 </p></td>\n
                <td><p align="center">{ProjectName}</p></td>\n
                <td><p align="center">子项目编号 </p></td>\n
                <td><p align="center">{ItemNo}</p></td>'''.format(ProjectName=ProjectName,ItemNo=ItemNo)
        else:
            projectTem = '''<td><p align="center">项目名称 </p></td>\n*\s*<td><p align="center">.*</p></td>\n*\s*<td><p align="center">合同编号 </p></td>\n*\s*<td><p align="center">.*</p></td>'''
            replaceProjectTem = '''<td><p align="center">项目名称 </p></td>\n
                <td><p align="center">{ProjectName}</p></td>\n
                <td><p align="center">合同编号 </p></td>\n
                <td><p align="center">{ItemNo}</p></td>'''.format(ProjectName=ProjectName,ItemNo=ItemNo)
        soup = re.sub(projectTem,replaceProjectTem,soup)


        # 4 设置 客户姓名 和 单位
        customerTem= '''<td><p align="center">客户姓名 </p></td>\n*\s*<td><p align="center">.*</p></td>\n*\s*<td><p align="center">客户单位 </p></td>\n*\s*<td><p align="center">.*</p></td>'''
        replaceCustomerTem = '''<td><p align="center">客户姓名 </p></td>\n
			<td><p align="center">{CustomerName}</p></td>\n
			<td><p align="center">客户单位 </p></td>\n
			<td><p align="center">{UnitAddress}</p></td>'''.format(CustomerName=CustomerName,UnitAddress=UnitAddress)
        soup = re.sub(customerTem,replaceCustomerTem,soup)

        # 5 设置 收样人 和 收样日期 (收样日期根据我们的收样单进行修改)
        ReceivedTem = '''<td><p align="center">收样人 </p></td>\n*\s*<td><p align="center">.*</p></td>\n*\s*<td><p align="center">收样日期 </p></td>\n*\s*<td><p align="center">.*</p></td>'''
        replaceReceivedTem = '''<td><p align="center">收样人 </p></td>\n
			<td><p align="center">{Received}</p></td>\n
			<td><p align="center">收样日期 </p></td>\n
			<td><p align="center">{ReceivedDate}</p></td>'''.format(Received=Received,ReceivedDate=ReceivedDate)
        soup = re.sub(ReceivedTem,replaceReceivedTem,soup)

        # 6 设置 审核日期 改成 与 报告日期同一天
        auditTem ='''<td><p align="center">审核日期 </p></td>\n*\s*<td><p align="center">.*</p></td>'''
        replaceAuditTem = '''<td><p align="center">审核日期 </p></td>\n<td><p align="center">{AuditDate}</p></td>'''.format(AuditDate=AuditDate)
        self.soup = re.sub(auditTem,replaceAuditTem,soup)

    # 设置页眉页脚
    def replaceHeardFooter(self,path):
        # 获取* docx
        docxsFiles = glob.glob(path+"/*.docx")
        for docxFile in docxsFiles:
            file_Name = docxFile.strip(".docx")
            # 解析文件
            docs = docx.Document(docxFile)
            # 获取时间戳
            ticks = time.time()
            ticks_pdf = file_Name + str(ticks) +".pdf"
            # docx to PDF 
            self.wordToPdf(docxFile,ticks_pdf) 
            # 获取页数
            reader = PdfFileReader(ticks_pdf)
            page_num = reader.getNumPages()
            
            # 删除不用的 pdf
            os.remove(ticks_pdf)
            # 循环读取word页 替换数据 删除数据
            for i in range(page_num):
                # # 替换文本
                for j, para in enumerate(docs.paragraphs):
                    text = para.text
                    dateText =  re.search("\w{4}/\d{0,2}/\d{0,2}\s\d{0,2}:\d{0,2}:\d{0,2}",text)
                    dateText2 = re.search("\w{4}-\d{0,2}-\d{0,2}_\d{0,2}-\d{0,2}-\d{0,2}.xad",text)
                    if text.strip() == "Assay Class: Data Path:":
                        para.text = ""
                    if text.strip() == "Plant RNA Nano":
                        para.text = ""
                    if text.strip() == "Created: Modified:":
                        para.text = ""
                    if dateText:
                        para.text = ""
                    if dateText2:
                        para.text = ""
            
                # 删除页眉页脚
                header = docs.sections[i].header # 获取第一个节的页眉
                paragraph_header = header.paragraphs[0]     # 获取页眉的第一个段落
                paragraph_header.text = "" # 添加页脚内容

                footer = docs.sections[i].footer # 获取第一个节的页脚
                paragraph_footer = footer.paragraphs[0] # 获取页脚的第一个段落
                paragraph_footer.text = "" # 添加页脚内容

            #保存文档
            docs.save(docxFile)

            # del pdf          
            outFile = "%s.pdf"%file_Name
            if os.path.exists(outFile):
                os.remove(outFile)

            # word TO PDF
            self.wordToPdf(docxFile,outFile)  
            
            os.remove(docxFile)


# Personal Novo 真核质检和无参质检 共通化方法
class TranscriptomeQuality(BaseHTML):

    def __init__(self,htmlPath,jsonPath):
        # 调用父类
        super().__init__(jsonPath)

        self.htmlPath = htmlPath

        with open(self.htmlPath,"r",encoding="utf-8")as f:
            self.soup = f.read()

        # self.soup = BeautifulSoup(htmlfile,"lxml")

    def main(self,anapath,parentDir):
        try:
            # #=================== 一、文件替换  ===================
            # --  1.src\Caliper 中PDF 替换成无logo版：
            oldlogoPath =  anapath +"/src/Caliper/Agilent 5400 峰图说明-RNA.pdf"
            if os.path.exists(oldlogoPath):
                os.remove(oldlogoPath)
                
                logoPath = parentDir + "/replaceFiles/NovoGene/无参质检/Agilent 5400 峰图说明-RNA.pdf"
                copyfile(logoPath ,oldlogoPath)  

            # -- 2.  src\images中 Logo替换
            oldlogoPath =  anapath +"/src/images/logo.png"
            if os.path.exists(oldlogoPath):
                os.remove(oldlogoPath)
            logoPath = parentDir + "/replaceFiles/logo.png"
            copyfile(logoPath ,oldlogoPath)

            #=================== 二、网页修改 ===================

            # 2.替换
            #   --（1）派森诺生物科技股份有限公司 替换为 苏州帕诺米克生物医药科技有限公司
            #   --（2）派森诺 替换为 帕诺米克。
            replaceDatas = self.jsonData.get("replaceFields")
            for replaceData in replaceDatas:
                name = replaceData.get("key")
                replaceName = replaceData.get("value")
                self.replaceText(name,replaceName)

            #项目基础信息修改
            self.editProinfo()

            # 写入文件
            self.writeFile()

            return SUCCESS_FLAG
        except Exception as e:
            print(e)
            return FAILED_FLAG