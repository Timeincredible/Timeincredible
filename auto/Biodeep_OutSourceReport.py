## -*- coding: utf-8 -*-

import glob
import docx
import  os, re
import comtypes.client
from pdf2docx import parse
from docx.shared import Pt
from docx.shared import Cm, Inches
from PyPDF2 import PdfFileReader, PdfFileWriter
# from win32com.client import Dispatch, constants, gencache

#######################################################################################
# 定义替换内容
#######################################################################################
SUCCESS_FLAG = "OK"
FAILED_FLAG = "NG"

replace_dict = {
   #"www.novogene.com":"http://www.bionovogene.com/",
   "国内 NGS 技术服务中心宏基因业务线":"帕诺米克 技术服务中心宏基因业务线",
   "诺禾致源": "帕诺米克",
   "诺禾": "诺米",
   "致源": "诺米",
}
delete_dict = [
   "北京诺禾致源科技股份有限公司 	网址：www.novogene.com 	邮箱：service@novogene.com 	电话：010-8283 7801",
]

HEADER_CONTENT = "苏州帕诺米克生物医药科技有限公司 网址:http://www.bionovogene.com/ 邮箱：project@bionovogene.com 电话：0512-62959105"
FOOTER_CONTENT = "苏州帕诺米克生物医药科技有限公司 网址:http://www.bionovogene.com/ 邮箱：project@bionovogene.com 电话：0512-62959105"


#######################################################################################
#
# 定义页脚内容
#
#######################################################################################
def del_files(path):
   fileNames = glob.glob(path + r'/*')

   for fileName in fileNames:
      try:
         os.remove(fileName)
      except:
         os.rmdir(fileName)
   os.rmdir(path)

######################################################################################
#
# 分割pdf文件，并转换为word文件
#
######################################################################################
# 1、原pdf文件，按单页，存单页pdf文件
def split_pdf(readFileName,analysis_dir):

   # 获取一个 PdfFileReader 对象
   pdfReader = PdfFileReader(open(readFileName, 'rb'))
   
   # 获取 PDF 的页数
   pageCount = pdfReader.getNumPages()
   pages = 0
   while pages < pageCount:
      
      # 返回一个 PageObject
      page = pdfReader.getPage(pages)
      
      # 获取一个 PdfFileWriter 对象
      pdfWriter = PdfFileWriter()
      
      # 将一个 PageObject 加入到 PdfFileWriter 中
      pdfWriter.addPage(page)
      
      writeFileName = analysis_dir + "/" + str(pages) + ".pdf"
      wordFileName = analysis_dir + "/" + str(pages) + ".docx"
      # 输出到文件中
      pdfWriter.write(open(writeFileName, 'wb'))
      
      #转成word文件
      parse(writeFileName, wordFileName)

      pages += 1
      
   return (pages)

######################################################################################
#
# word 文件合并
#
######################################################################################   
def combine_word_documents(merge_file, files):
   
   merged_document = docx.Document()

   for index, file in enumerate(files):
      sub_doc = docx.Document(file)


      # Don't add a page break if you've reached the last file.
      if index < len(files) - 1:
         sub_doc.add_page_break()

      for element in sub_doc.element.body:
         merged_document.element.body.append(element)

   merged_document.save(merge_file)

######################################################################################
#
# 删除文本中特定字段值
#
######################################################################################
def delete_paragraph(paragraph):
   p = paragraph._element
   p.getparent().remove(p)
   # p._p = p._element = None
   paragraph._p = paragraph._element = None

def del_content(document, delete_dict):

   for i, para in enumerate(document.paragraphs):
      for del_con in delete_dict:
         if del_con in para.text:
            delete_paragraph(para)
            break

   # 删除分页行
   for i, para in enumerate(document.paragraphs):
      match_result = re.match("^\d+$", para.text, flags=0)
      if match_result:
         delete_paragraph(para)

   return document

######################################################################################
#
# 替换文本中特定字段值
#
######################################################################################
def check_and_change(document, replace_dict):
   """
   遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。 
   （key 和 value 都是replace_dict中的键值对。）
   """
   
   #for i, para in enumerate(document.paragraphs):
   for para in document.paragraphs:
      for key, value in replace_dict.items():
         if key in para.text:
            para.text = para.text.replace(key, value)

   #for para in document.paragraphs:
      #line_text = ""
      #for i in range(len(para.runs)):
         ##print (objprint(para.runs[i]))
         #for key, value in replace_dict.items():
            #if key in para.runs[i].text:
               ##print(key+"->"+value)
               #para.runs[i].text = para.runs[i].text.replace(key, value)
               
   return document

######################################################################################
#
# 替换logo图标
#
######################################################################################
def change_logo(document):
   
   for section in document.sections:
      
      section.header.is_linked_to_previous = True
      section.footer.is_linked_to_previous = True
      
      foot = section.footer
      foot.paragraphs[0].text = FOOTER_CONTENT
      
      #footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      #section.font.size = Pt(15) 
      
   document.styles['Normal'].font.size = Pt(10.5)


   #print('图形图像的数量：', len(document.inline_shapes))
   
   if document.paragraphs[0].text.strip() == "":
      document.paragraphs[0].text= ""
   
   for i, para in enumerate(document.paragraphs):
      if i > 3:
         if para.text.strip() == "":
            para.text = ""
   
   
   #dict_rel = document.part._rels
   #for rel in dict_rel:
      #rel = dict_rel[rel]
      #if "image" in rel.target_ref:
         #print (rel.target_ref)
 
      
   return document
 


######################################################################################
#
# 诺禾 16s 分析报告 OTUanalysis 章节处理
#
######################################################################################
def OTUanalysisReport(number_files, analysis_dir):
   files = []
   
   for j in range(number_files):
      
      document = docx.Document(analysis_dir + "/" + str(j) + ".docx")
      
      # 处理顶部logo
      for i, para in enumerate(document.paragraphs):
         if para.text.strip() == "":
            para.text = ""
 
      # 处理页脚
      for i, para in enumerate(document.paragraphs):
         for del_con in delete_dict:
            if del_con in para.text:
               delete_paragraph(para)
               break

      # 删除分页行
      for i, para in enumerate(document.paragraphs):
         match_result = re.match("^\d+$", para.text, flags=0)
         if match_result:
            delete_paragraph(para)

      # 表格处理 读表格
      for tb in document.tables:
         #tb.autofit = False
         
         # for r in tb.rows:
         #   print (r.cells[0].text)
         #   print (r.cells[1].text)
         #   print(r.cells[2].text)
        
         col = tb.columns[0] 
         col.width = Inches(1)

         col1 = tb.columns[1] 
         col1.width = Inches(1.2)
         
         col2 = tb.columns[2] 
         col2.width = Inches(1)

      document = check_and_change(document, replace_dict)
      document.save(analysis_dir + "/" + str(j) + "_2.docx")
      files.append(analysis_dir + "/" + str(j) + "_2.docx")

   combine_word_documents(analysis_dir + "/result.docx", files)

def WebShowanalysisReport(number_files, analysis_dir):
   
   files = []
   
   for j in range(number_files):
      
      document = docx.Document(analysis_dir + "/" + str(j) + ".docx")
      
      # 处理顶部logo
      for i, para in enumerate(document.paragraphs):
         if para.text.strip() == "":
            para.text = ""
 
      # 处理页脚
      for i, para in enumerate(document.paragraphs):
         for del_con in delete_dict:
            if del_con in para.text:
               delete_paragraph(para)
               break

      # 删除分页行
      for i, para in enumerate(document.paragraphs):
         match_result = re.match("^\d+$", para.text, flags=0)
         if match_result:
            delete_paragraph(para)

      document = check_and_change(document, replace_dict)
      document.save(analysis_dir + "/" + str(j) + "_2.docx")
      files.append(analysis_dir + "/" + str(j) + "_2.docx")

   combine_word_documents(analysis_dir + "/result.docx", files)

######################################################################################
#
# 7/13 wwt msg: word TO pdf 
#
######################################################################################
def wordToPdf(sourceFile ,outFile):

   word = comtypes.client.CreateObject("Word.Application")
   word.Visible = 0
   newpdf = word.Documents.Open(sourceFile)

   try:
      if os.path.exists(outFile):
         os.remove(outFile)
      with open(outFile, 'w')as f: 
         f.close()
      newpdf.SaveAs(outFile, FileFormat=17)

      newpdf.Close()
      return SUCCESS_FLAG

   except Exception as e:
      newpdf.Close()
      return FAILED_FLAG


######################################################################################
#
# 7/13 wwt msg: 获取文件目录下的pdf文件 @return:list
#
######################################################################################
def getPdfFile(project_dir):
   '''获取pdf路径'''
   pdfFiles = glob.glob( project_dir +"/[0-9]*.*.pdf")
   return pdfFiles

######################################################################################
#
# 7/13 wwt msg: 设置页眉页脚  
#
######################################################################################
def setHeaderFooter(number_files,file):
   '''设置页眉页脚'''
   for i in range(number_files):
      
      docs = docx.Document(file)
      header = docs.sections[i].header # 获取第一个节的页眉
   
      paragraph = header.paragraphs[0]     # 获取页眉的第一个段落
      paragraph.text =  HEADER_CONTENT # 添加页面内容

      footer = docs.sections[i].footer # 获取第一个节的页脚
      paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
      paragraph.text = FOOTER_CONTENT # 添加页脚内容

   # 保存文档
   docs.save(file) 

   return SUCCESS_FLAG

######################################################################################
#
# 7/13 wwt msg: 共通化流程
#
######################################################################################
def setPublicMain(pdfFile):
   nowPath = os.getcwd()

   fileGroup = re.search("\d+\.\w+\.pdf",pdfFile)

   fileName = fileGroup.group()

   # step-2 : 定义临时数据文件存放目录
   analysis_dir = nowPath + "/tmp/{fileName}".format(fileName=fileName.strip(".pdf"))
   # 创建|删除文件
   if os.path.exists(analysis_dir):
      del_files(analysis_dir)
   os.makedirs(analysis_dir)

   # step-3 : 分割成单页
   number_files = split_pdf(pdfFile,analysis_dir)

   # step-4 : 分割成单页
   OTUanalysisReport(number_files, analysis_dir)

   # step-5 : word 转换成新的pdf文件
   newFilePath = pdfFile.replace(fileName,"") 
   
   # step-6: 设置页眉页脚
   wordPath = nowPath + "/tmp/{fileName}/result.docx".format(fileName=fileName.strip(".pdf"))
   setHeaderFooter(number_files,wordPath)

   # step-7 : word to  pdf 
   nowPath = os.getcwd()
   flag = wordToPdf(wordPath,newFilePath + "home.pdf")
  
   return flag

#SET  02.OTUanalysis
class setOTUanalysis():
   def __init__(self,proPath):
      self.proPath = proPath 

   def main(self):
      try:
         setPublicMain(self.proPath)
         return SUCCESS_FLAG
      except Exception as e:
         print(e)
         return FAILED_FLAG

#SET  03.AlphaDiversity
class setAlphaDiversity():
   def __init__(self,proPath):
      self.proPath = proPath

   def main(self):
      try:
         setPublicMain(self.proPath)
         return SUCCESS_FLAG
      except Exception as e:
         print(e)
         return FAILED_FLAG

#SET  04.BetaDiversity
class setBetaDiversity():
   def __init__(self,proPath):
      self.proPath = proPath

   def main(self):
      try:
         setPublicMain(self.proPath)
         return SUCCESS_FLAG
      except Exception as e:
         print(e)
         return FAILED_FLAG

#SET 05.Webshow
class setWebshow():
   def __init__(self,proPath):
      self.proPath = proPath

   def main(self):
      try:
        
         nowPath = os.getcwd()

         fileGroup = re.search("\w+\.pdf",self.proPath)
         fileName = fileGroup.group()

         # step-2 : 定义临时数据文件存放目录
         analysis_dir = nowPath + "/tmp/{fileName}".format(fileName=fileName.strip(".pdf"))
         # 创建|删除文件
         if os.path.exists(analysis_dir):
            del_files(analysis_dir)
         os.makedirs(analysis_dir)

         # step-3 : 分割成单页
         number_files = split_pdf(self.proPath,analysis_dir)

         # step-4 : 分割成单页
         WebShowanalysisReport(number_files, analysis_dir)

         # step-5 : word 转换成新的pdf文件
         newFilePath = self.proPath.replace(fileName,"") 

         # step-6: 设置页眉页脚
         wordPath = "./tmp/{fileName}/result.docx".format(fileName=fileName.strip(".pdf"))
         setHeaderFooter(number_files,wordPath)

         # step-7 : word to  pdf 
         nowPath = os.getcwd()
         flag = wordToPdf(nowPath + wordPath,newFilePath + "home.pdf")
      
         return SUCCESS_FLAG
      except Exception as e:
         print(e)
         return FAILED_FLAG

#SET 06.FunnctionPrediction
class setFunnctionPrediction():
   def __init__(self,proPath):
      self.proPath = proPath

   def main(self):
      try:
         setPublicMain(self.proPath)
         return SUCCESS_FLAG
      except Exception as e:
         print(e)
         return FAILED_FLAG

if __name__ == "__main__":
   pass
